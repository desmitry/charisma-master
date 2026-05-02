"""Tests for LLMClient provider dispatch in ml_worker."""

from __future__ import annotations

from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from charisma_schemas import AnalyzeProvider, PersonaRoles, SpeechReport


class TestLLMClientConstants:
    def test_transcription_limit(self, llm_client_module):
        assert llm_client_module.LLMClient.TRANSCRIPTION_LIMIT == 7_000

    def test_presentation_text_limit(self, llm_client_module):
        assert llm_client_module.LLMClient.PRESENTATION_TEXT_LIMIT == 5_000


class TestLLMClientInstantiation:
    def test_instantiates_with_patched_clients(self, llm_client_module):
        with (
            patch(
                "app.logic.llm_client.openai.AsyncOpenAI",
                return_value=MagicMock(name="openai_client"),
            ) as mock_openai,
            patch(
                "app.logic.llm_client.GigaChat",
                return_value=MagicMock(name="gigachat_client"),
            ) as mock_gigachat,
        ):
            client = llm_client_module.LLMClient()

        assert client is not None
        mock_openai.assert_called_once()
        mock_gigachat.assert_called_once()


class TestAnalyzeSpeech:
    def _valid_speech_report_json(self) -> str:
        return (
            '{"summary": "good", "structure": "clear", '
            '"mistakes": "none", "ideal_text": "great", '
            '"persona_feedback": "encouraging", '
            '"dynamic_fillers": [], '
            '"presentation_feedback": "nice", '
            '"competition_analysis": "market context", '
            '"useful_links": "link"}'
        )

    async def test_analyze_speech_gigachat_path(self, llm_client_module):
        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
            patch(
                "app.logic.llm_client.prompts.get_analyze_speech_system_prompt",
                return_value="system prompt",
            ),
        ):
            client = llm_client_module.LLMClient()
            client._call_gigachat = AsyncMock(
                return_value=self._valid_speech_report_json()
            )
            client._call_openai = AsyncMock()

            result = await client.analyze_speech(
                transcript_text="hi there",
                presentation_text="slides text",
                provider=AnalyzeProvider.gigachat,
                persona=PersonaRoles.speech_review_specialist,
            )

        assert isinstance(result, SpeechReport)
        assert result.summary == "good"
        client._call_gigachat.assert_awaited_once()
        client._call_openai.assert_not_awaited()

    async def test_analyze_speech_openai_path(self, llm_client_module):
        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
            patch(
                "app.logic.llm_client.prompts.get_analyze_speech_system_prompt",
                return_value="system prompt",
            ),
        ):
            client = llm_client_module.LLMClient()
            client._call_openai = AsyncMock(
                return_value=self._valid_speech_report_json()
            )
            client._call_gigachat = AsyncMock()

            result = await client.analyze_speech(
                transcript_text="hi",
                presentation_text="slides",
                provider=AnalyzeProvider.openai,
                persona=PersonaRoles.strict_critic,
            )

        assert isinstance(result, SpeechReport)
        assert result.summary == "good"
        client._call_openai.assert_awaited_once()
        client._call_gigachat.assert_not_awaited()

    async def test_long_transcript_is_truncated(self, llm_client_module):
        """Transcript longer than TRANSCRIPTION_LIMIT gets truncated."""
        long_text = "x" * 10_000
        presentation_text = "y" * 10_000

        captured_messages = {}

        async def capture_call(messages, model):
            captured_messages["messages"] = messages
            captured_messages["model"] = model
            return (
                '{"summary": "ok", "structure": "", '
                '"mistakes": "", "ideal_text": "", '
                '"persona_feedback": "", "dynamic_fillers": [], '
                '"presentation_feedback": "", "competition_analysis": "", '
                '"useful_links": ""}'
            )

        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
            patch(
                "app.logic.llm_client.prompts.get_analyze_speech_system_prompt",
                return_value="sys",
            ),
        ):
            client = llm_client_module.LLMClient()
            client._call_gigachat = AsyncMock(side_effect=capture_call)

            await client.analyze_speech(
                transcript_text=long_text,
                presentation_text=presentation_text,
                provider=AnalyzeProvider.gigachat,
                persona=PersonaRoles.kind_mentor,
            )

        user_msg = captured_messages["messages"][1]["content"]
        # TRANSCRIPTION_LIMIT=7000, PRESENTATION_TEXT_LIMIT=5000
        # So we expect at most 7000 x's in the transcription section
        # and at most 5000 y's in the presentation section.
        assert user_msg.count("x") <= 7_000
        assert user_msg.count("y") <= 5_000
        assert user_msg.count("x") == 7_000
        assert user_msg.count("y") == 5_000
        assert "ТРАНСКРИПЦИЯ" in user_msg
        assert "ТЕКСТ ПРЕЗЕНТАЦИИ" in user_msg
        assert "ИССЛЕДОВАНИЕ КОНКУРЕНТОВ" in user_msg

    async def test_analyze_speech_includes_competition_context(
        self, llm_client_module
    ):
        captured_messages = {}

        async def capture_call(messages, model):
            captured_messages["messages"] = messages
            captured_messages["model"] = model
            return self._valid_speech_report_json()

        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
            patch(
                "app.logic.llm_client.prompts.get_analyze_speech_system_prompt",
                return_value="sys",
            ),
        ):
            client = llm_client_module.LLMClient()
            client._call_openai = AsyncMock(side_effect=capture_call)

            result = await client.analyze_speech(
                transcript_text="hi",
                presentation_text="slides",
                provider=AnalyzeProvider.openai,
                persona=PersonaRoles.strict_critic,
                competition_analysis="Competitor A is stronger in distribution.",
            )

        assert result.competition_analysis == "market context"
        user_msg = captured_messages["messages"][1]["content"]
        assert "ИССЛЕДОВАНИЕ КОНКУРЕНТОВ" in user_msg
        assert "Competitor A is stronger in distribution." in user_msg

    async def test_gigachat_not_initialized_returns_error_report(
        self, llm_client_module
    ):
        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat", return_value=None),
            patch(
                "app.logic.llm_client.prompts.get_analyze_speech_system_prompt",
                return_value="sys",
            ),
        ):
            client = llm_client_module.LLMClient()
            # gigachat_client will be None; _call_gigachat raises ValueError
            result = await client.analyze_speech(
                transcript_text="hi",
                presentation_text="",
                provider=AnalyzeProvider.gigachat,
                persona=PersonaRoles.steve_jobs_style,
            )

        assert isinstance(result, SpeechReport)
        # Error flow returns generic error message
        assert (
            result.summary
            == "Analysis could not be completed. Please try again."
        )

    async def test_call_gigachat_api_error_propagates(
        self,
        llm_client_module,
    ):
        """GigaChat API exception propagates upward from _call_gigachat."""
        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
        ):
            client = llm_client_module.LLMClient()
            client.gigachat_client.achat = AsyncMock(
                side_effect=RuntimeError("GigaChat rate limit exceeded"),
            )

            with pytest.raises(RuntimeError, match="rate limit"):
                await client._call_gigachat(
                    messages=[{"role": "user", "content": "test"}],
                    model="GigaChat",
                )


class TestParseJsonResponse:
    def test_strips_markdown_fences(self, llm_client_module):
        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
        ):
            client = llm_client_module.LLMClient()

        raw = '```json\n{"summary": "x"}\n```'
        result = client._parse_json_response(raw)
        assert result == {"summary": "x"}

    def test_invalid_json_raises_json_decode_error(self, llm_client_module):
        """Invalid JSON raises RuntimeError with JSON decode error details."""
        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
        ):
            client = llm_client_module.LLMClient()

        with pytest.raises(RuntimeError, match="LLM returned invalid JSON"):
            client._parse_json_response("not json")


class TestGetEvaluationCriteria:
    @pytest.mark.asyncio
    async def test_from_json_preset(self, llm_client_module, tmp_path):
        """JSON preset file yields list[EvaluationCriterion]."""
        criteria_file = tmp_path / "criteria.json"
        criteria_file.write_text(
            '{"criteria": ['
            '{"name": "Clarity", "description": "c", "max_value": 10},'
            '{"name": "Structure", "description": "s", "max_value": 5}'
            "]}"
        )

        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
        ):
            client = llm_client_module.LLMClient()
            result = await client.get_evaluation_criteria(str(criteria_file))

        assert len(result) == 2
        assert result[0].name == "Clarity"
        assert result[0].max_value == 10
        assert result[0].current_value == 0
        assert result[0].feedback == ""

    @pytest.mark.asyncio
    async def test_empty_json_raises_runtime_error(
        self,
        llm_client_module,
        tmp_path,
    ):
        """Empty JSON preset raises RuntimeError."""
        criteria_file = tmp_path / "empty.json"
        criteria_file.write_text('{"criteria": []}')

        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
        ):
            client = llm_client_module.LLMClient()

            with pytest.raises(RuntimeError):
                await client.get_evaluation_criteria(str(criteria_file))

    @pytest.mark.asyncio
    async def test_from_txt_uses_llm(self, llm_client_module, tmp_path):
        """TXT file routes through LLM to extract criteria."""
        criteria_file = tmp_path / "criteria.txt"
        criteria_file.write_text("1. Speak clearly\n2. Don't use fillers")

        fake_response = (
            '```json\n{"criteria": ['
            '{"name": "Clarity", "description": "Speak", "max_value": 10}'
            "]}\n```"
        )

        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
            patch(
                "app.logic.llm_client.prompts.get_evaluation_criteria_identity_prompt",
                return_value="sys",
            ),
        ):
            client = llm_client_module.LLMClient()
            client._call_gigachat = AsyncMock(return_value=fake_response)

            result = await client.get_evaluation_criteria(str(criteria_file))

        assert len(result) == 1
        assert result[0].name == "Clarity"

    @pytest.mark.asyncio
    async def test_unsupported_format_raises(self, llm_client_module):
        """Unsupported extension raises RuntimeError."""
        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
        ):
            client = llm_client_module.LLMClient()

            with pytest.raises(RuntimeError):
                await client.get_evaluation_criteria("/tmp/criteria.pdf")

    @pytest.mark.asyncio
    async def test_empty_llm_response_raises(
        self,
        llm_client_module,
        tmp_path,
    ):
        """LLM returned empty criteria list raises RuntimeError."""
        criteria_file = tmp_path / "criteria.txt"
        criteria_file.write_text("Some criteria")

        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
            patch(
                "app.logic.llm_client.prompts.get_evaluation_criteria_identity_prompt",
                return_value="sys",
            ),
        ):
            client = llm_client_module.LLMClient()
            client._call_gigachat = AsyncMock(return_value='{"criteria": []}')

            with pytest.raises(RuntimeError):
                await client.get_evaluation_criteria(str(criteria_file))


class TestAnalyzeWithEvaluationCriteria:
    @pytest.mark.asyncio
    async def test_happy_path(self, llm_client_module):
        """Normal scoring yields criteria with current_value and feedback."""
        from charisma_schemas import EvaluationCriterion

        captured_messages = {}
        criteria = [
            EvaluationCriterion(
                name="Clarity",
                description="d",
                max_value=10,
            ),
            EvaluationCriterion(
                name="Structure",
                description="d",
                max_value=5,
            ),
        ]

        fake_response = (
            '{"criteria": ['
            '{"name": "Clarity", "current_value": 8, "feedback": "Good"},'
            '{"name": "Structure", "current_value": 4, "feedback": "OK"}'
            "]}"
        )

        async def capture_call(messages, model):
            captured_messages["messages"] = messages
            captured_messages["model"] = model
            return fake_response

        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
            patch(
                "app.logic.llm_client.prompts.get_evaluation_criteria_rate_prompt",
                return_value="sys",
            ),
        ):
            client = llm_client_module.LLMClient()
            client._call_gigachat = AsyncMock(side_effect=capture_call)

            result = await client.analyze_with_evalution_criteria(
                transcript_text="speech text",
                presentation_text="slides",
                provider=AnalyzeProvider.gigachat,
                evaluation_criteria=criteria,
                competition_analysis="Competitor B targets the same segment.",
            )

        assert len(result) == 2
        assert result[0].current_value == 8
        assert result[0].feedback == "Good"
        assert result[1].current_value == 4
        assert result[1].feedback == "OK"
        user_msg = captured_messages["messages"][1]["content"]
        assert "ИССЛЕДОВАНИЕ КОНКУРЕНТОВ" in user_msg
        assert "Competitor B targets the same segment." in user_msg

    @pytest.mark.asyncio
    async def test_identify_competition_subject_fills_default_query(
        self, llm_client_module
    ):
        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
            patch(
                "app.logic.llm_client.prompts.get_competition_product_prompt",
                return_value="sys",
            ),
        ):
            client = llm_client_module.LLMClient()
            client._call_openai = AsyncMock(
                return_value=(
                    '{"can_identify": true, "product_name": "Acme CRM", '
                    '"product_description": "CRM for small sales teams", '
                    '"search_query": ""}'
                )
            )

            result = await client.identify_competition_subject(
                transcript_text="text",
                presentation_text="slides",
                provider=AnalyzeProvider.openai,
            )

        assert result["can_identify"] is True
        assert result["product_name"] == "Acme CRM"
        assert result["search_query"] == "Acme CRM конкуренты"

    @pytest.mark.asyncio
    async def test_empty_list_returns_empty(self, llm_client_module):
        """Empty criteria list returns [] without calling the LLM."""
        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
        ):
            client = llm_client_module.LLMClient()
            result = await client.analyze_with_evalution_criteria(
                transcript_text="text",
                presentation_text="",
                provider=AnalyzeProvider.gigachat,
                evaluation_criteria=[],
            )

        assert result == []

    @pytest.mark.asyncio
    async def test_count_mismatch_raises(self, llm_client_module):
        """LLM returned a different number of criteria than requested."""
        from charisma_schemas import EvaluationCriterion

        criteria = [
            EvaluationCriterion(name="A", description="d", max_value=10),
            EvaluationCriterion(name="B", description="d", max_value=5),
        ]

        fake_response = (
            '{"criteria": ['
            '  {"name": "A", "current_value": 5, "feedback": "ok"},'
            '  {"name": "B", "current_value": 3, "feedback": "ok"},'
            '  {"name": "C", "current_value": 7, "feedback": "unexpected"}'
            "]}"
        )

        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
            patch(
                "app.logic.llm_client.prompts.get_evaluation_criteria_rate_prompt",
                return_value="sys",
            ),
        ):
            client = llm_client_module.LLMClient()
            client._call_gigachat = AsyncMock(return_value=fake_response)

            with pytest.raises(RuntimeError):
                await client.analyze_with_evalution_criteria(
                    transcript_text="text",
                    presentation_text="",
                    provider=AnalyzeProvider.gigachat,
                    evaluation_criteria=criteria,
                )

    @pytest.mark.asyncio
    async def test_invalid_current_value_raises(self, llm_client_module):
        """LLM returned current_value=-1 raises RuntimeError."""
        from charisma_schemas import EvaluationCriterion

        criteria = [
            EvaluationCriterion(name="Clarity", description="d", max_value=10),
        ]

        fake_response = (
            '{"criteria": ['
            '  {"name": "Clarity", "current_value": -1, "feedback": "bad"}'
            "]}"
        )

        with (
            patch("app.logic.llm_client.openai.AsyncOpenAI"),
            patch("app.logic.llm_client.GigaChat"),
            patch(
                "app.logic.llm_client.prompts.get_evaluation_criteria_rate_prompt",
                return_value="sys",
            ),
        ):
            client = llm_client_module.LLMClient()
            client._call_gigachat = AsyncMock(return_value=fake_response)

            with pytest.raises(RuntimeError):
                await client.analyze_with_evalution_criteria(
                    transcript_text="text",
                    presentation_text="",
                    provider=AnalyzeProvider.gigachat,
                    evaluation_criteria=criteria,
                )


class TestReadDocumentText:
    def test_txt_file(self, llm_client_module, tmp_path):
        """TXT file returns full text."""
        txt = tmp_path / "doc.txt"
        txt.write_text("Line 1\nLine 2")

        result = llm_client_module.LLMClient._read_document_text(
            str(txt),
            ".txt",
        )
        assert result == "Line 1\nLine 2"

    def test_docx_file(self, llm_client_module, tmp_path):
        """DOCX file text extracted from paragraphs."""
        from docx import Document

        docx = tmp_path / "doc.docx"
        doc = Document()
        doc.add_paragraph("Hello world")
        doc.add_paragraph("Second paragraph")
        doc.save(str(docx))

        result = llm_client_module.LLMClient._read_document_text(
            str(docx),
            ".docx",
        )
        assert "Hello world" in result
        assert "Second paragraph" in result

    def test_empty_docx_raises_runtime_error(
        self,
        llm_client_module,
        tmp_path,
    ):
        """Empty DOCX raises RuntimeError."""
        from docx import Document

        docx = tmp_path / "empty.docx"
        Document().save(str(docx))

        with pytest.raises(RuntimeError):
            llm_client_module.LLMClient._read_document_text(
                str(docx),
                ".docx",
            )

    def test_unsupported_format_raises(self, llm_client_module):
        """PDF extension raises RuntimeError."""
        with pytest.raises(RuntimeError):
            llm_client_module.LLMClient._read_document_text(
                "/tmp/doc.pdf",
                ".pdf",
            )
